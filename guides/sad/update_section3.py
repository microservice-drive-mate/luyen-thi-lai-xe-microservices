"""
Script to update Section 3 "Architectural Representation" in ADD_v2.docx
to match the actual codebase implementation.

Run from repo root:
    python guides/sad/update_section3.py
"""

import sys
import os
from docx import Document
from docx.shared import Pt

DOCX_PATH = os.path.join(os.path.dirname(__file__), "ADD_v2.docx")


# ---------------------------------------------------------------------------
# Content definitions
# ---------------------------------------------------------------------------

SECTION3_INTRO = (
    "DriveMate adopts a Microservices architecture following the 4+1 View Model, "
    "with Kong as the API Gateway, Keycloak as the identity provider, and Consul for "
    "per-service configuration management. Each microservice is a stateless NestJS "
    "application owning its own PostgreSQL database. Synchronous communication uses "
    "REST/HTTP for latency-sensitive operations; asynchronous communication uses "
    "RabbitMQ for background events (notifications, user sync, analytics)."
)

# ---- 3.1 Logical View ----

LOGICAL_VIEW_DESC = (
    "The Logical View presents the system's decomposition into functional microservices "
    "and their primary responsibilities. It shows how client applications interact with "
    "Kong API Gateway, which routes requests to the appropriate microservices. Kong "
    "validates JWT tokens via Keycloak on every request. Microservices communicate "
    "synchronously via HTTP for latency-sensitive operations (e.g. exam-service fetching "
    "questions from question-service), and asynchronously via RabbitMQ for background "
    "tasks (notifications, analytics, user sync). Each service reads its runtime "
    "configuration from Consul KV. Centralized logging is provided by the ELK stack "
    "(Elasticsearch, Logstash, Kibana) — all services ship logs via Winston HTTP transport."
)

LOGICAL_VIEW_SERVICES = [
    "Kong API Gateway — single entry point; rate limiting, CORS, declarative routing (DB-less mode).",
    "Keycloak 24.0 — JWT issuance, RBAC enforcement, SSO, realm management.",
    "identity-service — Keycloak admin API integration, token blacklist (Redis), forgot-password flow.",
    "user-service — user profiles, student details, license tier assignment, role sync via RabbitMQ.",
    "exam-service — exam templates (config), exam sessions, atomic scoring, question snapshot.",
    "question-service — question bank CRUD, soft-delete, version history; images in object storage.",
    "course-service — courses, lessons, enrollment, lesson progress.",
    "simulation-service — driving scenarios, server-side FSM for action validation.",
    "analytics-service — pre-aggregated progress stats, spaced repetition algorithm.",
    "notification-service — RabbitMQ consumer, email/push delivery.",
    "media-service — file metadata storage; actual files in Cloudflare R2 / Azure Blob (S3-compatible).",
    "docs-service — centralized Swagger UI aggregation.",
    "Consul KV — per-service configuration with priority: env vars > Consul > defaults.",
    "ELK Stack — centralized logging (Winston → Logstash → Elasticsearch → Kibana).",
]

LOGICAL_VIEW_PLANTUML = """\
@startuml DriveMate_LogicalView
skinparam defaultFontName Arial
skinparam defaultFontSize 12
skinparam backgroundColor #FFFFFF
skinparam component {
  BackgroundColor #D6E4F0
  BorderColor #1F4E79
  FontColor #1F4E79
}
skinparam rectangle {
  BackgroundColor #EEF3FB
  BorderColor #1F4E79
  FontColor #1F4E79
}
skinparam arrow {
  Color #1F4E79
}

actor Student
actor "Admin / Manager\\nInstructor" as Admin

rectangle "Client Layer" {
  component [Web App\\n(React / Vue)] as Web
}

rectangle "Gateway Layer" {
  component [Kong API Gateway\\n[Rate Limit · CORS · Routing]] as GW
}

rectangle "Identity Layer" {
  component [Keycloak 24.0\\n[JWT · RBAC · SSO]] as KC
}

rectangle "Config Layer" {
  component [Consul KV 1.19\\n[Per-service config]] as Consul
}

rectangle "Microservices Layer" {
  component [identity-service\\n(token blacklist, forgot-pwd)] as Auth
  component [user-service\\n(profiles, license tiers)] as User
  component [exam-service\\n(sessions, scoring, config)] as Exam
  component [question-service\\n(bank, versions)] as QB
  component [course-service\\n(courses, enrollment)] as Course
  component [simulation-service\\n(driving scenarios, FSM)] as DP
  component [analytics-service\\n(stats, SRS)] as SRS
  component [notification-service\\n(email, push)] as Notif
  component [media-service\\n(file metadata)] as Media
  component [docs-service\\n(Swagger aggregation)] as Docs
}

rectangle "Infrastructure Layer" {
  database "PostgreSQL 15\\n(9 isolated databases)" as PG
  database "Redis 7\\n(token blacklist)" as Redis
  queue "RabbitMQ 3\\n(durable event queues)" as MQ
  storage "Cloudflare R2 / Azure Blob\\n(S3-compatible)" as S3
}

rectangle "Observability Layer" {
  component [ELK Stack\\n(Elasticsearch · Logstash · Kibana)] as ELK
}

Student --> Web
Admin --> Web
Web --> GW : HTTPS/REST

GW --> KC : JWT validation
GW --> Auth
GW --> User
GW --> Exam
GW --> QB
GW --> Course
GW --> DP
GW --> SRS
GW --> Media

Auth --> KC : admin API
Auth --> PG
Auth --> Redis : token blacklist
User --> PG
User --> MQ : publish events
Exam --> PG
Exam --> MQ : publish events
Exam --> QB : HTTP (question pool)
QB --> PG
QB --> S3
Course --> PG
DP --> PG
SRS --> PG
SRS --> MQ : consume
Notif --> MQ : consume
Media --> PG
Media --> S3

Consul ..> Auth : config
Consul ..> User : config
Consul ..> Exam : config
Consul ..> QB : config

Auth ..> ELK : logs (Winston)
User ..> ELK : logs (Winston)
Exam ..> ELK : logs (Winston)

@enduml"""

# ---- 3.2 Implementation View ----

IMPL_VIEW_DESC = (
    "The Implementation View outlines how services are organized as code artifacts and "
    "how they communicate at the implementation level. Each NestJS microservice is an "
    "independent TypeScript application in a Turborepo monorepo, deployed as a Docker "
    "container. Authentication and authorization are handled by Keycloak guards injected "
    "at the module level — no individual service implements its own auth logic. All "
    "services share the @repo/common library for DDD base classes, Consul config loading, "
    "HTTP response/error helpers, and centralized Winston logging. Prisma v7 generates an "
    "isolated client package per service (e.g. @prisma/exam-client) to enforce the "
    "database-per-service boundary."
)

IMPL_VIEW_STACK = [
    "Backend Microservices: Node.js / NestJS 11 (TypeScript), Turborepo monorepo",
    "Shared Library: @repo/common — DDD base classes (AggregateRoot, Entity, ValueObject), ConsulConfigFactory, AppLoggerModule (Winston + ELK), ApiResponseInterceptor, DomainExceptionFilter",
    "API Gateway: Kong (declarative DB-less, kong.yaml / kong.dev.yaml)",
    "Identity Provider: Keycloak 24.0 (JWT, RBAC, @nestjs/keycloak-connect guards)",
    "Primary Databases: PostgreSQL 15 — 9 isolated databases, one per service",
    "ORM: Prisma v7 with @prisma/adapter-pg; custom client output per service",
    "Cache: Redis 7 — token blacklist (identity-service only)",
    "Message Broker: RabbitMQ 3 — durable queues, noAck: false",
    "Config Management: Consul 1.19 — KV store, per-service namespacing",
    "Object Storage: Cloudflare R2 / Azure Blob (S3-compatible, via media-service)",
    "Container Orchestration: Docker + Docker Compose",
    "Logging/Observability: Winston → Logstash (HTTP) → Elasticsearch → Kibana",
    "Code Quality: Biome (lint/format), TypeScript strict mode",
    "API Documentation: OpenAPI / Swagger (auto-generated per service, aggregated by docs-service)",
]

IMPL_VIEW_PLANTUML = """\
@startuml DriveMate_ImplementationView
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam backgroundColor #FFFFFF
skinparam package {
  BackgroundColor #EEF3FB
  BorderColor #1F4E79
  FontColor #1F4E79
}
skinparam component {
  BackgroundColor #D6E4F0
  BorderColor #1F4E79
}
skinparam database {
  BackgroundColor #FFF3CD
  BorderColor #856404
}
skinparam queue {
  BackgroundColor #D4EDDA
  BorderColor #155724
}
skinparam storage {
  BackgroundColor #F8D7DA
  BorderColor #721C24
}
skinparam arrow {
  Color #1F4E79
}

package "Client Application" {
  [Web App\\n(React / Vue)] as App
}

package "Gateway" {
  [Kong API Gateway\\n(DB-less, kong.yaml)] as GW
}

package "Identity & Config" {
  [Keycloak 24.0\\n(JWT, RBAC, realm)] as KC
  [Consul KV 1.19\\n(per-service config)] as Consul
}

package "Backend Microservices (NestJS / TypeScript)" {
  [identity-service] as Auth
  [user-service] as User
  [exam-service] as Exam
  [question-service] as QB
  [course-service] as Course
  [simulation-service] as DP
  [analytics-service] as SRS
  [notification-service] as Notif
  [media-service] as Media
  [docs-service] as Docs
}

package "@repo/common (Shared Library)" {
  [DDD Base Classes\\n(AggregateRoot, Entity, VO)] as DDD
  [ConsulConfigFactory\\n(config loader + Joi)] as CConf
  [AppLoggerModule\\n(Winston + ELK HTTP)] as Log
  [ApiResponseInterceptor\\n(unified response shape)] as ARI
  [DomainExceptionFilter\\n(domain to HTTP error)] as DEF
}

package "Data & Messaging Infrastructure" {
  database "PostgreSQL 15\\n(9 isolated DBs\\nvia Prisma custom clients)" as PG
  database "Redis 7\\n(token blacklist)" as Redis
  queue "RabbitMQ 3\\n(durable queues, noAck:false)" as MQ
  storage "Cloudflare R2 / Azure Blob\\n(S3-compatible)" as S3
}

package "Observability (ELK)" {
  [Logstash :5044] as Logstash
  database "Elasticsearch :9200" as ES
  [Kibana :5601] as Kibana
}

App --> GW : REST/HTTPS

GW --> KC : JWT validation
GW --> Auth
GW --> User
GW --> Exam
GW --> QB
GW --> Course
GW --> DP
GW --> SRS
GW --> Media

Auth ..> KC : admin API
Auth ..> DDD
Auth ..> CConf
Auth ..> Log
Auth ..> ARI
Auth ..> DEF

User ..> DDD
Exam ..> DDD
QB ..> DDD

Exam --> MQ : publish
SRS --> MQ : consume
Notif --> MQ : consume
User --> MQ : publish

Auth --> PG
Auth --> Redis
Exam --> PG
SRS --> PG
QB --> PG
QB --> S3
Course --> PG
DP --> PG
User --> PG
Media --> PG
Media --> S3

Consul ..> Auth : load config
Consul ..> User : load config
Consul ..> Exam : load config

Auth ..> Logstash : logs
User ..> Logstash : logs
Exam ..> Logstash : logs
Logstash --> ES
ES --> Kibana

@enduml"""

# ---- 3.3 Deployment View ----

DEPLOY_VIEW_DESC = (
    "The Deployment View describes the physical distribution of system components. "
    "Two deployment configurations are documented: the development stack (Docker Compose) "
    "used during active development, and the production target (Kubernetes) for scalable "
    "cloud deployment."
)

DEPLOY_DEV_DESC = (
    "The development stack runs entirely via Docker Compose (docker-compose.yaml for full "
    "stack, docker-compose.infra.yml for hybrid mode where services run on the host machine "
    "and infrastructure runs in Docker). Kong uses kong.dev.yaml and routes to "
    "host.docker.internal:{port} when services run locally. Consul is seeded at startup "
    "via the consul-init container using consul-seed-development.json. Mailpit replaces a "
    "real SMTP server for email testing."
)

DEPLOY_DEV_PLANTUML = """\
@startuml DriveMate_DeploymentView_Dev
skinparam defaultFontName Arial
skinparam defaultFontSize 10
skinparam backgroundColor #FFFFFF
skinparam node {
  BackgroundColor #EEF3FB
  BorderColor #1F4E79
  FontColor #1F4E79
}
skinparam component {
  BackgroundColor #D6E4F0
  BorderColor #1F4E79
}
skinparam database {
  BackgroundColor #FFF3CD
  BorderColor #856404
}
skinparam queue {
  BackgroundColor #D4EDDA
  BorderColor #155724
}
skinparam storage {
  BackgroundColor #F8D7DA
  BorderColor #721C24
}
skinparam arrow {
  Color #1F4E79
}

title DriveMate — Development Deployment (Docker Compose)

node "Host Machine / Docker Compose Stack" {

  node "Gateway" {
    component [Kong\\n(DB-less, kong.dev.yaml)\\nPort: 8000] as Kong
  }

  node "Application Services" {
    component [identity-service :3001] as Auth
    component [user-service :3002] as User
    component [exam-service :3003] as Exam
    component [course-service :3004] as Course
    component [question-service :3005] as QB
    component [notification-service :3006] as Notif
    component [analytics-service :3007] as SRS
    component [simulation-service :3008] as DP
    component [docs-service :3009] as Docs
    component [media-service :3010] as Media
  }

  node "Identity & Config" {
    component [Keycloak :8080] as KC
    component [Consul :8500\\n(consul-init seeds KV)] as Consul
  }

  node "Data Services" {
    database "PostgreSQL x9\\n(:5432 - :5440)" as PG
    database "Redis :6379\\n(token blacklist)" as Redis
    queue "RabbitMQ :5672\\n(Management :15672)" as MQ
  }

  node "Observability (ELK)" {
    component [Logstash :5044] as Logstash
    database "Elasticsearch :9200" as ES
    component [Kibana :5601] as Kibana
  }

  node "Dev Tools" {
    component [Mailpit :1025 / :8025\\n(SMTP mock)] as Mailpit
  }

}

Kong --> Auth
Kong --> User
Kong --> Exam
Kong --> QB
Kong --> Course
Kong --> DP
Kong --> SRS
Kong --> Media

Auth --> KC
Auth --> PG
Auth --> Redis
User --> PG
User --> MQ
Exam --> PG
Exam --> MQ
Exam --> QB : HTTP
QB --> PG
Course --> PG
DP --> PG
SRS --> PG
SRS --> MQ
Notif --> MQ
Notif --> Mailpit : SMTP
Media --> PG

Consul ..> Auth : config
Consul ..> User : config
Consul ..> Exam : config

Auth ..> Logstash : Winston HTTP
User ..> Logstash : Winston HTTP
Logstash --> ES --> Kibana

@enduml"""

DEPLOY_PROD_DESC = (
    "For production deployment, all NestJS microservices run as stateless Kubernetes "
    "Deployments with Horizontal Pod Autoscaling (HPA) to handle peak exam traffic. "
    "External traffic enters through an Nginx Ingress Controller that terminates TLS "
    "before forwarding to Kong. Keycloak and Consul are deployed as StatefulSets with "
    "persistent volumes. Managed data services (PostgreSQL x9, Redis HA cluster, "
    "RabbitMQ cluster with persistent queues) are provisioned independently of the "
    "application layer. The ELK stack (Elasticsearch, Logstash, Kibana) provides "
    "centralized log aggregation and visualization."
)

DEPLOY_PROD_PLANTUML = """\
@startuml DriveMate_DeploymentView_Prod
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam backgroundColor #FFFFFF
skinparam node {
  BackgroundColor #EEF3FB
  BorderColor #1F4E79
  FontColor #1F4E79
}
skinparam component {
  BackgroundColor #D6E4F0
  BorderColor #1F4E79
}
skinparam database {
  BackgroundColor #FFF3CD
  BorderColor #856404
}
skinparam queue {
  BackgroundColor #D4EDDA
  BorderColor #155724
}
skinparam storage {
  BackgroundColor #F8D7DA
  BorderColor #721C24
}
skinparam arrow {
  Color #1F4E79
}

node "Cloud Provider (AWS / GCP / Azure)" {

  node "Kubernetes Cluster" {
    node "Ingress Layer" {
      component [Nginx Ingress Controller\\n(TLS Termination)] as Ingress
      component [Kong\\n(DB-less gateway)] as GW
    }
    node "Application Pods (HPA-Scaled)" {
      component [identity-service Pod(s)] as Auth
      component [user-service Pod(s)] as User
      component [exam-service Pod(s)] as Exam
      component [question-service Pod(s)] as QB
      component [course-service Pod(s)] as Course
      component [simulation-service Pod(s)] as DP
      component [analytics-service Pod(s)] as SRS
      component [notification-service Pod(s)] as Notif
      component [media-service Pod(s)] as Media
      component [docs-service Pod(s)] as Docs
    }
    node "Identity & Config (StatefulSets)" {
      component [Keycloak\\n(JWT, RBAC)] as KC
      component [Consul\\n(KV config)] as Consul
    }
    node "Observability Stack (ELK)" {
      component [Logstash] as Logstash
      component [Kibana] as Kibana
    }
  }

  node "Managed Data Services" {
    database "PostgreSQL x9\\n(Primary + Read Replica)" as PG
    database "Redis Cluster (HA)" as Redis
    queue "RabbitMQ Cluster\\nPersistent Queues" as MQ
    storage "S3-Compatible\\nObject Storage" as S3
    database "Elasticsearch\\n(Log storage)" as ES
  }
}

Ingress --> GW : route
GW --> KC : JWT validation
GW --> Auth
GW --> User
GW --> Exam
GW --> QB
GW --> Course
GW --> DP
GW --> SRS
GW --> Media

Auth --> KC
Auth --> PG
Auth --> Redis
Exam --> PG
Exam --> MQ
User --> PG
User --> MQ
SRS --> MQ
Notif --> MQ
QB --> PG
QB --> S3
Course --> PG
DP --> PG
Media --> PG
Media --> S3

Consul ..> Auth : config
Consul ..> User : config
Consul ..> Exam : config

Auth ..> Logstash : logs
Exam ..> Logstash : logs
Logstash --> ES --> Kibana

@enduml"""

# ---- 3.4 Data View ----

DATA_VIEW_DESC = (
    "The Data View describes the core entities, their relationships, and the data storage "
    "strategy. DriveMate applies the database-per-service pattern: each microservice owns "
    "an isolated PostgreSQL 15 database with no cross-service foreign keys. Services "
    "reference each other only by UUID. Eventual consistency is maintained via domain "
    "events published through RabbitMQ."
)

DATA_VIEW_TABLE_HEADER = "Database allocation per service:"

DATA_VIEW_DB_ROWS = [
    ("identity-service", "identity_db :5432", "IdentityUser, TokenAudit"),
    ("user-service", "user_db :5433", "UserProfile, StudentDetail, LicenseAssignmentAudit"),
    ("exam-service", "exam_db :5434", "ExamTemplate, ExamSession, ExamSessionQuestion (snapshot)"),
    ("question-service", "question_db :5436", "Question, QuestionOption, QuestionTopic"),
    ("course-service", "course_db :5435", "Course, Lesson, Enrollment, LessonProgress"),
    ("simulation-service", "simulation_db :5439", "SimulationScenario, SimulationSession"),
    ("notification-service", "notification_db :5437", "NotificationRecord"),
    ("analytics-service", "analytics_db :5438", "ProgressStat, SRSItem"),
    ("media-service", "media_db :5440", "MediaFile"),
]

DATA_VIEW_REDIS = (
    "Caching Strategy (Redis — identity-service only):\n"
    "Token blacklist: blacklist:{jti} — TTL = token remaining lifetime; checked by "
    "TokenBlacklistGuard on every authenticated request across all services."
)

DATA_VIEW_QUEUES = (
    "Message Queues (RabbitMQ):\n"
    "user_service_events — consumed by user-service (e.g. identity.user.created).\n"
    "exam_service_publish — published by exam-service.\n"
    "analytics_service_events — consumed by analytics-service.\n"
    "notification_service_events — consumed by notification-service.\n"
    "media_service_events — consumed by media-service."
)

DATA_VIEW_SCHEMA_NOTE = (
    "Key Schema Highlights:\n\n"
    "user_db: UserProfile {id (=Keycloak sub), fullName, email, role, isActive} | "
    "StudentDetail {licenseTier (A1–F), enrolledAt} | "
    "LicenseAssignmentAudit {oldTier, newTier, changedById, changedAt}\n\n"
    "exam_db: ExamTemplate {licenseCategory, totalQuestions, passingScore, durationMinutes} | "
    "ExamSession {studentId (cross-service ref), status, score, isPassed} | "
    "ExamSessionQuestion {questionContent (snapshot), optionsSnapshot (JSONB), isCorrect, isBookmarked}"
)

DATA_VIEW_PLANTUML = """\
@startuml DriveMate_DataView
skinparam defaultFontName Arial
skinparam defaultFontSize 10
skinparam backgroundColor #FFFFFF
skinparam entity {
  BackgroundColor #D6E4F0
  BorderColor #1F4E79
  FontColor #000000
}
skinparam rectangle {
  BackgroundColor #EEF3FB
  BorderColor #1F4E79
}
skinparam arrow {
  Color #1F4E79
}

title DriveMate - Data View (Database-per-Service)

rectangle "identity_db" {
  entity IdentityUser {
    * id : UUID <<PK>>
    --
    keycloakId : VARCHAR UNIQUE
    email : VARCHAR
    createdAt : TIMESTAMP
  }
}

rectangle "user_db" {
  entity UserProfile {
    * id : UUID <<PK>> (=Keycloak sub)
    --
    fullName : VARCHAR
    email : VARCHAR UNIQUE
    phoneNumber : VARCHAR
    role : ENUM[ADMIN|CENTER_MANAGER|INSTRUCTOR|STUDENT]
    isActive : BOOLEAN
  }
  entity StudentDetail {
    * id : UUID <<PK>>
    --
    userProfileId : UUID (ref)
    licenseTier : ENUM[A1|A2|B1|B2|C|D|E|F]
    enrolledAt : TIMESTAMP
    notes : TEXT
  }
  entity LicenseAssignmentAudit {
    * id : UUID <<PK>>
    --
    studentId : UUID (ref)
    oldTier : ENUM
    newTier : ENUM
    changedById : UUID (ref)
    changedAt : TIMESTAMP
  }
  UserProfile ||--o| StudentDetail : "has (if STUDENT)"
  UserProfile ||--o{ LicenseAssignmentAudit : "audited by"
}

rectangle "exam_db" {
  entity ExamTemplate {
    * id : UUID <<PK>>
    --
    licenseCategory : ENUM[A1..F]
    totalQuestions : INT
    passingScore : INT
    durationMinutes : INT
  }
  entity ExamSession {
    * id : UUID <<PK>>
    --
    studentId : UUID (cross-service ref)
    templateId : UUID <<FK>>
    status : ENUM[IN_PROGRESS|COMPLETED|TIMED_OUT|CANCELLED]
    score : INT
    isPassed : BOOLEAN
    startedAt : TIMESTAMP
    finishedAt : TIMESTAMP
  }
  entity ExamSessionQuestion {
    * id : UUID <<PK>>
    --
    sessionId : UUID <<FK>>
    questionId : UUID (cross-service ref)
    questionContent : TEXT (snapshot)
    optionsSnapshot : JSONB (snapshot)
    selectedOptionId : UUID
    isCorrect : BOOLEAN
    isBookmarked : BOOLEAN
    answeredAt : TIMESTAMP
  }
  ExamTemplate ||--o{ ExamSession : "used by"
  ExamSession ||--o{ ExamSessionQuestion : "contains"
}

rectangle "question_db" {
  entity Question {
    * id : UUID <<PK>>
    --
    licenseCategory : ENUM
    content : TEXT
    explanation : TEXT
    isActive : BOOLEAN
    isFatal : BOOLEAN
  }
  entity QuestionOption {
    * id : UUID <<PK>>
    --
    questionId : UUID <<FK>>
    content : TEXT
    isCorrect : BOOLEAN
    displayOrder : INT
  }
  Question ||--o{ QuestionOption : "has"
}

note bottom of ExamSessionQuestion
  Question content & options are SNAPSHOTTED
  at exam start. Future edits to question_db
  do NOT affect in-progress or completed exams.
end note

rectangle "Redis (identity-service)" {
  entity TokenBlacklist {
    key: "blacklist:{jti}"
    TTL: remaining token lifetime
    Used by: TokenBlacklistGuard
  }
}

rectangle "RabbitMQ Queues" {
  entity Queues {
    user_service_events
    exam_service_publish
    analytics_service_events
    notification_service_events
    media_service_events
  }
}

@enduml"""

# ---- 3.5 Process View ----

PROCESS_VIEW_DESC = (
    "The Process View describes how DriveMate handles concurrency, asynchronous tasks, "
    "offline resilience, and server-side state control at runtime. Four key process "
    "flows are documented below."
)

PROCESS_FLOW1 = (
    "Flow 1 — Login via Keycloak:\n"
    "When a user submits login credentials, the request travels through Kong to the "
    "identity-service. The identity-service forwards the credentials to the Keycloak "
    "token endpoint; Keycloak handles password hashing, account lockout, and session "
    "management internally. On success, a JWT and refresh token are returned to the "
    "client. On logout, the identity-service adds the JWT's jti claim to the Redis "
    "token blacklist with TTL equal to the token's remaining lifetime. All subsequent "
    "requests are checked against the blacklist by TokenBlacklistGuard, which is "
    "registered globally across all services."
)

PROCESS_FLOW2 = (
    "Flow 2 — Asynchronous Notification:\n"
    "When an instructor sends a study alert, the notification-service accepts the "
    "request and immediately returns HTTP 202 to the instructor's interface without "
    "blocking it. The notification task is published to a RabbitMQ queue and consumed "
    "in the background. Failed deliveries are retried automatically; messages that "
    "exhaust retries are moved to a Dead Letter Queue for inspection. This decoupling "
    "ensures that notification delivery never slows down the instructor's UI."
)

PROCESS_FLOW3 = (
    "Flow 3 — Offline Sync (Auto-Save):\n"
    "During an active exam, the client auto-saves answers to the exam-service every "
    "5 to 10 seconds using an idempotent upsert operation. If the network is lost, the "
    "client detects the disconnection within 2 seconds and shows a non-blocking offline "
    "indicator. Answers continue to be stored locally. When connectivity is restored, "
    "the client silently syncs the buffered answers using the same idempotent endpoint, "
    "ensuring no duplicates are created regardless of how many retries occur."
)

PROCESS_FLOW4 = (
    "Flow 4 — Server-Side FSM (Simulation Service):\n"
    "All state transition logic for driving practice resides exclusively on the server "
    "within simulation-service. The client sends a user action event and waits for the "
    "server's response before updating the UI. The simulation-service validates the "
    "action against the current FSM state and either advances the state or rejects the "
    "action with a structured error containing the error code and severity. The client "
    "renders the appropriate feedback within 100 milliseconds of receiving the response. "
    "No FSM logic is shipped to the client, preventing manipulation through tampered builds."
)

PROCESS_VIEW_PLANTUML = """\
@startuml DriveMate_ProcessView
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam backgroundColor #FFFFFF
skinparam sequence {
  ArrowColor #1F4E79
  ActorBorderColor #1F4E79
  LifeLineBorderColor #1F4E79
  ParticipantBorderColor #1F4E79
  ParticipantBackgroundColor #D6E4F0
  BoxBorderColor #1F4E79
}

title DriveMate - Process View (Key Flows)

== Flow 1: Login via Keycloak ==
actor Student
participant "Kong\\nAPI Gateway" as GW
participant "identity-service" as Auth
participant "Keycloak" as KC
database "Redis" as Redis
database "identity_db" as PG

Student -> GW : POST /auth/login (credentials)
GW -> Auth : Forward credentials
Auth -> KC : POST /token (Keycloak token endpoint)
KC -> KC : Verify credentials (BCrypt, lockout logic)
KC -> Auth : JWT + refresh_token
Auth -> PG : Write audit record
Auth -> GW : JWT + refresh_token
GW -> Student : Return tokens

note over KC: Account lockout and password\\nhashing handled by Keycloak

== On Logout ==
Student -> GW : POST /auth/logout
GW -> Auth : Forward logout
Auth -> Redis : SET blacklist:{jti} TTL=remaining
Auth -> GW : 200 OK
note over Auth,Redis: TokenBlacklistGuard on all\\nservices rejects blacklisted JWTs

== Flow 2: Asynchronous - Notification ==
participant "notification-service" as Notif
queue "RabbitMQ" as MQ

participant "Instructor Client" as Inst
Inst -> GW : Send study alert
GW -> Notif : Trigger alert (sync)
Notif -> MQ : Publish to notification_service_events
Notif -> Inst : HTTP 202 Accepted (immediate)
MQ -> Notif : Consume + deliver
note right: Retry on failure\\nDLQ after max retry

== Flow 3: Offline Sync - Auto-Save ==
participant "Client App" as App
participant "exam-service" as ESS

App -> App : Auto-save every 5-10 s (local storage)
App -[#red]x GW : Network lost
note over App: Offline indicator\\nshown within 2 s
App -> App : Buffer answers locally
...network restored...
App -> GW : Sync buffered answers (idempotent)
GW -> ESS : Upsert answers
ESS -> App : Sync confirmed

== Flow 4: Server-Side FSM - Simulation ==
participant "simulation-service" as DP

Student -> GW : Action event (e.g. shift gear)
GW -> DP : Forward action
DP -> DP : Validate against current FSM state
alt Valid sequence
  DP -> GW : next_valid_state + result
  GW -> Student : Render new state
else Invalid sequence
  DP -> GW : HTTP 400 + error_code + severity
  GW -> Student : Show colour-coded alert (<= 100 ms)
end

@enduml"""

# ---- 3.6 Scenario View ----

SCENARIO_VIEW_DESC = (
    "The Scenario View illustrates how components across multiple architectural layers "
    "collaborate to fulfil the four most critical use cases in DriveMate. Each scenario "
    "is selected because it exercises multiple quality attributes simultaneously."
)

SCENARIO1_DESC = (
    "This scenario covers the most critical transactional flow in the system. A student "
    "submits their completed exam. The exam-service opens a single database transaction "
    "that records all answers, grades the submission by comparing selected answers to "
    "stored correct answers, evaluates fatal questions first (a single wrong answer on a "
    "fatal question produces an immediate fail verdict), writes the immutable COMPLETED "
    "result, and updates SRS Item records — all in one atomic unit. If any step fails, "
    "the entire transaction rolls back and the student receives a safe-to-retry error. "
    "Once the transaction commits, the COMPLETED status is permanent; no role including "
    "Admin may modify it. This scenario exercises ASR-REL-04, ASR-DI-01, and ASR-DI-07."
)

SCENARIO1_PLANTUML = """\
@startuml DriveMate_Scenario1_AtomicSubmit
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam backgroundColor #FFFFFF
skinparam sequence {
  ArrowColor #1F4E79
  ParticipantBorderColor #1F4E79
  ParticipantBackgroundColor #D6E4F0
}

title Scenario 1: Atomic Exam Submission & Grading

actor Student
participant "Client App" as App
participant "Kong API Gateway" as GW
participant "exam-service" as ES
database "exam_db\\n(PostgreSQL)" as PG

Student -> App : Tap Submit
App -> App : Collect all answers + flag states
App -> GW : POST submit (answers + flags)
GW -> ES : Forward submission

ES -> PG : BEGIN TRANSACTION
ES -> PG : Write exam answers
ES -> PG : Grade answers (server-side)
note right: Check fatal questions first
ES -> PG : Write final result (COMPLETED + immutable)
ES -> PG : Update SRS_Item records
ES -> PG : COMMIT

alt Commit success
  ES -> GW : Score + pass/fail result
  GW -> App : Display result screen
else Any step fails
  ES -> PG : ROLLBACK (entire operation)
  ES -> GW : HTTP 500 - safe to retry
  GW -> App : Show retry prompt
end

note over PG: COMPLETED status is immutable\\nafter commit - no role can alter it

@enduml"""

SCENARIO2_DESC = (
    "This scenario covers the real-time driving simulation. When the student opens the "
    "practice screen, all audio command assets are downloaded and buffered before the "
    "session becomes active. During the session, each student action is sent to the "
    "server-side FSM in simulation-service for validation. Valid actions advance the "
    "state and trigger audio playback with under 100 milliseconds of latency. Invalid "
    "or out-of-sequence actions are rejected by the server and a colour-coded alert is "
    "rendered within 100 milliseconds of the response: red for fatal errors, yellow for "
    "point deductions. No state logic runs on the client. This scenario exercises "
    "ASR-SEC-07, ASR-UX-02, and ASR-UX-03."
)

SCENARIO2_PLANTUML = """\
@startuml DriveMate_Scenario2_DrivingPractice
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam backgroundColor #FFFFFF
skinparam sequence {
  ArrowColor #1F4E79
  ParticipantBorderColor #1F4E79
  ParticipantBackgroundColor #D6E4F0
}

title Scenario 2: Driving Practice Session (Server FSM + Audio)

actor Student
participant "Client App" as App
participant "Kong API Gateway" as GW
participant "simulation-service" as DP
storage "Object Storage\\n(Audio MP3)" as S3

== Session Start: Audio Preload ==
Student -> App : Open practice screen
App -> S3 : Preload ALL audio command assets
S3 -> App : Audio files buffered (ready)
note over App: Session locked until\\nall assets fully loaded

== Active Session: Action Loop ==
Student -> App : Perform driving action (e.g. press clutch)
App -> GW : Send action event
GW -> DP : Forward to server-side FSM

DP -> DP : Validate action against\\ncurrent FSM state

alt Valid sequence
  DP -> GW : next_valid_state
  GW -> App : Render new state
  App -> App : Play audio command\\n(< 100 ms latency)
else Invalid sequence
  DP -> GW : error_code + severity (FATAL / DEDUCT)
  GW -> App : Render colour-coded alert\\n(<= 100 ms from response)
  note over App: RED = fatal error\\nYELLOW = point deduction
end

note over DP: No FSM logic\\nexists on client

@enduml"""

SCENARIO3_DESC = (
    "This scenario covers exam resilience under network loss. During a normal exam "
    "session the client auto-saves answers to the exam-service every 5 to 10 seconds "
    "using an idempotent endpoint so that duplicate saves never create duplicate records. "
    "When the network is lost, the client detects the disconnection within 2 seconds and "
    "displays a non-blocking offline banner without covering the question content. The "
    "student continues answering and the client stores all new answers locally. When "
    "connectivity returns, the offline indicator clears automatically and the buffered "
    "answers are synced silently without any student action. On final submission, all "
    "answers — including those saved offline — are included in the atomic submission "
    "payload. This scenario exercises ASR-REL-03 and ASR-UX-05."
)

SCENARIO3_PLANTUML = """\
@startuml DriveMate_Scenario3_OfflineSync
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam backgroundColor #FFFFFF
skinparam sequence {
  ArrowColor #1F4E79
  ParticipantBorderColor #1F4E79
  ParticipantBackgroundColor #D6E4F0
}

title Scenario 3: Offline Exam & Auto-Sync

actor Student
participant "Client App" as App
participant "Kong API Gateway" as GW
participant "exam-service" as ESS

== Normal Operation ==
Student -> App : Answer questions
App -> GW : Auto-save (every 5-10 s, idempotent)
GW -> ESS : Upsert answers

== Network Loss ==
App -[#red]x GW : Connection lost
note over App: Detect loss within 2 s\\nShow offline indicator\\n(non-blocking)

Student -> App : Continue answering
App -> App : Buffer in local storage
App -> App : Auto-save locally (every 5-10 s)

== Network Restored ==
App -> App : Detect reconnection
App -> App : Clear offline indicator
App -> GW : Sync buffered answers (idempotent POST)
GW -> ESS : Upsert - duplicate saves ignored

note over ESS: Idempotency key prevents\\nduplicate records on retry

ESS -> GW : Sync ACK
GW -> App : Silent confirmation\\n(no user action needed)

== Submission ==
Student -> App : Tap Submit
App -> GW : Final submit with all answers
GW -> ESS : Atomic transaction (answers + result)

@enduml"""

SCENARIO4_DESC = (
    "This scenario covers exam creation and the enforcement of answer confidentiality. "
    "When a student starts an exam, the exam-service loads the active ExamTemplate and "
    "immediately snapshots the question content as immutable JSON blobs attached to each "
    "ExamSessionQuestion record — future question edits cannot affect this exam. Questions "
    "are selected using indexed database queries grouped by topic to match the configured "
    "counts exactly; full-table scans are prohibited. Before the question payload leaves "
    "the server, the Response Serialiser strips the isCorrect field from every option "
    "object. The entire generation flow completes within 3 seconds at the 95th percentile. "
    "This scenario exercises ASR-PERF-12, ASR-DI-09, and ASR-SEC-05."
)

SCENARIO4_PLANTUML = """\
@startuml DriveMate_Scenario4_ExamGen
skinparam defaultFontName Arial
skinparam defaultFontSize 11
skinparam backgroundColor #FFFFFF
skinparam sequence {
  ArrowColor #1F4E79
  ParticipantBorderColor #1F4E79
  ParticipantBackgroundColor #D6E4F0
}

title Scenario 4: Exam Generation & Answer Confidentiality

actor Student
participant "Client App" as App
participant "Kong API Gateway" as GW
participant "exam-service" as ES
participant "question-service" as QS
participant "Response\\nSerialiser" as RS
database "exam_db\\n(PostgreSQL)" as PG

Student -> App : Tap Start Exam
App -> GW : POST start exam
GW -> ES : Generate exam

ES -> PG : Load active ExamTemplate
ES -> QS : GET questions by category (HTTP)
note right: Indexed query, no full scan
QS -> ES : Question list (with options)
ES -> ES : Shuffle answer options\\n(correct identity preserved)
ES -> PG : Write ExamSession + ExamSessionQuestions\\n(snapshot content + options as JSONB)

alt Exam satisfies template config
  ES -> RS : Pass question payload
  RS -> RS : STRIP isCorrect field\\nfrom every option object
  RS -> GW : Clean payload (no correct answers)
  GW -> App : Exam ready - questions displayed
  note over App: P95 < 3 s\\nfrom request to first question
else Cannot satisfy config
  ES -> GW : HTTP 400 - descriptive error
  GW -> App : Show error message
end

note over RS: ApiResponseInterceptor\\nstrips sensitive fields\\nbefore every response

@enduml"""


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def remove_paragraphs_from(doc, start_index):
    """Remove all paragraphs and body-level elements from start_index onwards."""
    body = doc.element.body
    # Get ALL direct children of body (paragraphs + tables + sectPr, etc.)
    all_children = list(body)

    # Find the element at doc.paragraphs[start_index]
    target_elem = doc.paragraphs[start_index]._element

    # Find position of target_elem in body children
    cut_pos = None
    for i, child in enumerate(all_children):
        if child is target_elem:
            cut_pos = i
            break

    if cut_pos is None:
        print(f"WARNING: Could not find paragraph {start_index} in body children")
        return

    # Remove all children from cut_pos onwards (but NOT the last sectPr if present)
    children_to_remove = all_children[cut_pos:]
    for child in children_to_remove:
        # Keep the section properties element (w:sectPr) at the very end
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_normal(doc, text):
    doc.add_paragraph(text, style='Normal')


def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Paragraph')


def add_plantuml(doc, plantuml_text):
    """Add PlantUML code block — one Normal paragraph per line."""
    add_normal(doc, 'PlantUML Source:')
    add_normal(doc, '')
    for line in plantuml_text.split('\n'):
        add_normal(doc, line)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    total_paras = len(doc.paragraphs)
    print(f"Total paragraphs before edit: {total_paras}")

    # Find section 3 heading paragraph index
    sec3_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '3. Architectural Representation' and p.style.name == 'Heading 1':
            sec3_idx = i
            break

    if sec3_idx is None:
        print("ERROR: Could not find '3. Architectural Representation' Heading 1")
        sys.exit(1)

    print(f"Found Section 3 heading at paragraph index: {sec3_idx}")

    # Remove everything after the Section 3 heading (index sec3_idx+1 onwards)
    remove_paragraphs_from(doc, sec3_idx + 1)
    print(f"Paragraphs after removal: {len(doc.paragraphs)}")

    # ---- Insert new Section 3 content ----

    # Intro
    add_normal(doc, '')
    add_normal(doc, SECTION3_INTRO)
    add_normal(doc, '')

    # --- 3.1 Logical View ---
    add_h2(doc, '3.1 Logical View')
    add_normal(doc, '')
    add_normal(doc, LOGICAL_VIEW_DESC)
    add_normal(doc, '')
    add_normal(doc, 'Key services and responsibilities:')
    for svc in LOGICAL_VIEW_SERVICES:
        add_bullet(doc, svc)
    add_normal(doc, '')
    add_plantuml(doc, LOGICAL_VIEW_PLANTUML)

    # --- 3.2 Implementation View ---
    add_normal(doc, '')
    add_h2(doc, '3.2 Implementation View')
    add_normal(doc, '')
    add_normal(doc, IMPL_VIEW_DESC)
    add_normal(doc, '')
    add_normal(doc, 'Technology stack:')
    for item in IMPL_VIEW_STACK:
        add_bullet(doc, item)
    add_normal(doc, '')
    add_plantuml(doc, IMPL_VIEW_PLANTUML)

    # --- 3.3 Deployment View ---
    add_normal(doc, '')
    add_h2(doc, '3.3 Deployment View')
    add_normal(doc, '')
    add_normal(doc, DEPLOY_VIEW_DESC)
    add_normal(doc, '')

    add_h3(doc, '3.3.1 Development Deployment (Docker Compose)')
    add_normal(doc, '')
    add_normal(doc, DEPLOY_DEV_DESC)
    add_normal(doc, '')
    add_plantuml(doc, DEPLOY_DEV_PLANTUML)
    add_normal(doc, '')

    add_h3(doc, '3.3.2 Production Target Deployment (Kubernetes)')
    add_normal(doc, '')
    add_normal(doc, DEPLOY_PROD_DESC)
    add_normal(doc, '')
    add_plantuml(doc, DEPLOY_PROD_PLANTUML)

    # --- 3.4 Data View ---
    add_normal(doc, '')
    add_h2(doc, '3.4 Data View')
    add_normal(doc, '')
    add_normal(doc, DATA_VIEW_DESC)
    add_normal(doc, '')
    add_normal(doc, DATA_VIEW_TABLE_HEADER)
    add_normal(doc, '')
    for svc, db, entities in DATA_VIEW_DB_ROWS:
        add_bullet(doc, f'{svc} → {db}: {entities}')
    add_normal(doc, '')
    add_normal(doc, DATA_VIEW_REDIS)
    add_normal(doc, '')
    add_normal(doc, DATA_VIEW_QUEUES)
    add_normal(doc, '')
    add_normal(doc, DATA_VIEW_SCHEMA_NOTE)
    add_normal(doc, '')
    add_plantuml(doc, DATA_VIEW_PLANTUML)

    # --- 3.5 Process View ---
    add_normal(doc, '')
    add_h2(doc, '3.5 Process View')
    add_normal(doc, '')
    add_normal(doc, PROCESS_VIEW_DESC)
    add_normal(doc, '')
    add_normal(doc, PROCESS_FLOW1)
    add_normal(doc, '')
    add_normal(doc, PROCESS_FLOW2)
    add_normal(doc, '')
    add_normal(doc, PROCESS_FLOW3)
    add_normal(doc, '')
    add_normal(doc, PROCESS_FLOW4)
    add_normal(doc, '')
    add_plantuml(doc, PROCESS_VIEW_PLANTUML)

    # --- 3.6 Scenario View ---
    add_normal(doc, '')
    add_h2(doc, '3.6 Scenario View')
    add_normal(doc, '')
    add_normal(doc, SCENARIO_VIEW_DESC)
    add_normal(doc, '')

    add_h3(doc, '3.6.1 Scenario 1: Atomic Exam Submission & Grading')
    add_normal(doc, '')
    add_normal(doc, SCENARIO1_DESC)
    add_normal(doc, '')
    add_plantuml(doc, SCENARIO1_PLANTUML)
    add_normal(doc, '')

    add_h3(doc, '3.6.2 Scenario 2: Driving Practice Session (Server FSM + Audio)')
    add_normal(doc, '')
    add_normal(doc, SCENARIO2_DESC)
    add_normal(doc, '')
    add_plantuml(doc, SCENARIO2_PLANTUML)
    add_normal(doc, '')

    add_h3(doc, '3.6.3 Scenario 3: Offline Exam & Auto-Sync')
    add_normal(doc, '')
    add_normal(doc, SCENARIO3_DESC)
    add_normal(doc, '')
    add_plantuml(doc, SCENARIO3_PLANTUML)
    add_normal(doc, '')

    add_h3(doc, '3.6.4 Scenario 4: Exam Generation & Answer Confidentiality')
    add_normal(doc, '')
    add_normal(doc, SCENARIO4_DESC)
    add_normal(doc, '')
    add_plantuml(doc, SCENARIO4_PLANTUML)

    # Save
    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")
    print(f"Total paragraphs after edit: {len(doc.paragraphs)}")
    print("Done.")


if __name__ == '__main__':
    main()
